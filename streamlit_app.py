import streamlit as st
from llama_index.core import VectorStoreIndex, Document, Settings
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core.retrievers import VectorIndexRetriever
from llama_index.core.query_engine import RetrieverQueryEngine
from llama_index.core.postprocessor import SimilarityPostprocessor
from llama_index.embeddings.openai import OpenAIEmbedding
from googleapiclient.discovery import build
from google.oauth2 import service_account
import os
from io import BytesIO
import pandas as pd
from docx import Document as DocxDocument
from llama_index.core.storage import StorageContext
import pickle
import hashlib
import json
from pathlib import Path
import logging
import openai  # Import OpenAI SDK
from googlesearch import search
import requests
from bs4 import BeautifulSoup

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Set your OpenAI API key from Streamlit secrets
openai.api_key = st.secrets["OPEN_AI_KEY"]

# Load service account credentials from Streamlit secrets
service_account_info = json.loads(st.secrets["gcp"]["service_account_json"])

# Create credentials using the service account info
creds = service_account.Credentials.from_service_account_info(service_account_info)

# Build the Google Drive service
service = build('drive', 'v3', credentials=creds)

FOLDER_ID = '1H6PgbGSvDlTvc-Zip3VW0XiHjedDKrR9'

@st.cache_resource
def get_drive_service():
    return build('drive', 'v3', credentials=creds, cache_discovery=False)

def get_file_content(service, file_id, mime_type):
    try:
        if mime_type == 'application/vnd.google-apps.document':
            content = service.files().export(
                fileId=file_id,
                mimeType='text/plain'
            ).execute()
            return content.decode('utf-8')
        
        elif mime_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
            # Handle XLSX files
            content = service.files().get_media(fileId=file_id).execute()
            excel_file = BytesIO(content)
            
            # Read all sheets
            all_sheets = pd.read_excel(excel_file, sheet_name=None)
            formatted_content = []
            
            # Process each sheet
            for sheet_name, df in all_sheets.items():
                # Clean up the dataframe
                df = df.fillna('')  # Replace NaN with empty string
                
                # Add sheet name as context
                formatted_content.append(f"\nSheet: {sheet_name}\n")
                
                # Convert column names to string and make them lowercase for better searching
                df.columns = df.columns.astype(str).str.lower()
                
                # Process each row with column headers as context
                for idx, row in df.iterrows():
                    row_content = []
                    for col in df.columns:
                        if row[col] != '':  # Only include non-empty cells
                            row_content.append(f"{col}: {row[col]}")
                    formatted_content.append(" | ".join(row_content))
        
            return "\n".join(formatted_content)
            
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            # Handle DOCX files
            content = service.files().get_media(fileId=file_id).execute()
            doc_file = BytesIO(content)
            doc = DocxDocument(doc_file)
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            
        elif mime_type == 'text/plain':
            content = service.files().get_media(fileId=file_id).execute()
            return content.decode('utf-8')
            
        elif mime_type == 'application/vnd.google-apps.folder':
            # Skip folders
            return None
            
        else:
            return None
            
    except Exception as e:
        st.write(f"Error reading file {file_id}: {str(e)}")
        return None

def get_folder_contents(service, folder_id, documents=None):
    if documents is None:
        documents = []
    
    query = f"'{folder_id}' in parents"
    results = service.files().list(
        q=query,
        spaces='drive',
        fields='files(id, name, mimeType, webViewLink)',
        supportsAllDrives=True,
        includeItemsFromAllDrives=True
    ).execute()
    
    files = results.get('files', [])
    
    for file in files:
        if file['mimeType'] == 'application/vnd.google-apps.folder':
            # Recursively get contents of subfolder
            get_folder_contents(service, file['id'], documents)
        else:
            content = get_file_content(service, file['id'], file['mimeType'])
            if content:
                doc = Document(
                    text=content,
                    metadata={
                        "filename": file['name'],
                        "link": file.get('webViewLink', ''),
                        "filetype": file['mimeType']
                    }
                )
                documents.append(doc)
    
    return documents

def get_files_hash():
    """Get a hash of all files and their last modified times"""
    service = get_drive_service()
    files = []
    
    def collect_files(folder_id):
        query = f"'{folder_id}' in parents"
        results = service.files().list(
            q=query,
            fields='files(id, name, mimeType, modifiedTime)',
            supportsAllDrives=True,
            includeItemsFromAllDrives=True
        ).execute()
        
        for file in results.get('files', []):
            if file['mimeType'] == 'application/vnd.google-apps.folder':
                collect_files(file['id'])
            else:
                files.append({
                    'id': file['id'],
                    'modified': file['modifiedTime']
                })
    
    collect_files(FOLDER_ID)
    return hashlib.md5(json.dumps(files, sort_keys=True).encode()).hexdigest()

def save_index(index, doc_states):
    """Save index and document states to persistent storage"""
    # Create storage directory if it doesn't exist
    storage_dir = Path("storage")
    storage_dir.mkdir(exist_ok=True)
    
    storage_context = index._storage_context
    
    # Save all index components
    try:
        with open(storage_dir / 'docstore.pkl', 'wb') as f:
            pickle.dump(storage_context.docstore, f)
        with open(storage_dir / 'index_store.pkl', 'wb') as f:
            pickle.dump(storage_context.index_store, f)
        with open(storage_dir / 'vector_store.pkl', 'wb') as f:
            pickle.dump(storage_context.vector_store, f)
        with open(storage_dir / 'doc_states.json', 'w') as f:
            json.dump(doc_states, f)
            
        # Save current hash
        current_hash = get_files_hash()
        with open(storage_dir / 'files_hash.txt', 'w') as f:
            f.write(current_hash)
            
        return True
    except Exception as e:
        st.error(f"Error saving index: {str(e)}")
        return False

def load_index():
    """Load index from persistent storage"""
    storage_dir = Path("storage")
    
    try:
        # Verify hash hasn't changed
        current_hash = get_files_hash()
        hash_path = storage_dir / 'files_hash.txt'
        
        if hash_path.exists():
            stored_hash = hash_path.read_text().strip()
            if current_hash != stored_hash:
                logger.info("Hash mismatch. Rebuilding index.")
                return None
                
        # Load all components
        with open(storage_dir / 'docstore.pkl', 'rb') as f:
            docstore = pickle.load(f)
        with open(storage_dir / 'index_store.pkl', 'rb') as f:
            index_store = pickle.load(f)
        with open(storage_dir / 'vector_store.pkl', 'rb') as f:
            vector_store = pickle.load(f)
            
        # Create storage context
        storage_context = StorageContext.from_defaults(
            docstore=docstore,
            index_store=index_store,
            vector_store=vector_store
        )
        
        # Reconstruct index
        return VectorStoreIndex(
            nodes=list(docstore.docs.values()),
            storage_context=storage_context,
            show_progress=True
        )
        
    except FileNotFoundError:
        return None
    except Exception as e:
        st.error(f"Error loading index: {str(e)}")
        return None

@st.cache_resource
def create_index():
    """Create or load index with improved document processing"""
    try:
        # Try loading existing index first
        index = load_index()
        if index is not None:
            logger.info("Using existing index")
            return index
        
        # If no index exists, create new one
        logger.info("Creating new index...")
        documents = get_drive_documents()
        
        # Process documents with improved chunking
        processed_docs = []
        for doc in documents:
            chunks = get_document_chunks(doc)
            if isinstance(chunks, str):
                doc = Document(text=chunks, metadata=doc.metadata)
            else:
                doc = Document(text="\n\n".join(chunks), metadata=doc.metadata)
            processed_docs.append(doc)
        
        # Create new index
        index = VectorStoreIndex.from_documents(
            processed_docs,
            show_progress=True
        )
        
        # Save index and document states
        doc_states = get_document_states()
        if save_index(index, doc_states):
            logger.info("Index created and saved successfully")
        
        return index
        
    except Exception as e:
        st.error(f"Error creating index: {str(e)}")
        return None

def generate_response(index, user_input):
    if not index:
        return {
            "answer": "No documents found in the drive to search through.",
            "primary_citations": [],
            "secondary_citations": []
        }
    
    # Pre-prompt for semantic search
    pre_prompt = "You are searching for a journalistic topic. Please provide relevant details."
    user_input = f"{pre_prompt} {user_input}"

    retriever = VectorIndexRetriever(index=index, similarity_top_k=7)
    postprocessor = SimilarityPostprocessor(similarity_cutoff=0.77)
    query_engine = RetrieverQueryEngine.from_args(
        retriever=retriever, node_postprocessors=[postprocessor], response_mode="tree_summarize"
    )

    try:
        response = query_engine.query(user_input)
        if not response or not response.source_nodes:
            return {
                "answer": "I couldn't find any relevant information based on your query.",
                "primary_citations": [],
                "secondary_citations": []
            }
        
        source_docs = sorted(response.source_nodes, key=lambda x: getattr(x, 'score', 0), reverse=True)
        primary_citations = [doc for doc in source_docs if getattr(doc, 'score', 0) > 0.85][:2]
        secondary_citations = [doc for doc in source_docs if 0.77 <= getattr(doc, 'score', 0) <= 0.85][:3]
        
        # Use the updated OpenAI call to generate a response
        openai_response = generate_openai_response(user_input)
        
        formatted_response = {
            "answer": openai_response,
            "primary_citations": [format_citation(doc) for doc in primary_citations],
            "secondary_citations": [format_citation(doc) for doc in secondary_citations]
        }
        return formatted_response

    except Exception as e:
        return {
            "answer": f"An error occurred while processing your request: {str(e)}",
            "primary_citations": [],
            "secondary_citations": []
        }

def format_citation(doc):
    metadata = doc.metadata
    score = getattr(doc, 'score', None)
    snippet = doc.node.text[:150] + "..." if len(doc.node.text) > 150 else doc.node.text
    return {
        "filename": metadata.get('filename', 'Untitled'),
        "score": f"{score:.2f}" if score else 'N/A',
        "snippet": snippet,
        "link": metadata.get('link', '#')
    }

def get_drive_documents():
    service = get_drive_service()
    logger.info("Fetching files from Google Drive...")
    documents = get_folder_contents(service, FOLDER_ID)
    logger.info(f"Successfully created {len(documents)} document objects")
    return documents

def get_document_states():
    """Get current state of all documents in Drive"""
    service = get_drive_service()
    doc_states = {}
    
    def collect_file_states(folder_id):
        query = f"'{folder_id}' in parents"
        results = service.files().list(
            q=query,
            fields='files(id, name, mimeType, modifiedTime)',
            supportsAllDrives=True,
            includeItemsFromAllDrives=True
        ).execute()
        
        for file in results.get('files', []):
            if file['mimeType'] == 'application/vnd.google-apps.folder':
                collect_file_states(file['id'])
            else:
                doc_states[file['id']] = {
                    'name': file['name'],
                    'modified': file['modifiedTime']
                }
    
    collect_file_states(FOLDER_ID)
    return doc_states

def process_spreadsheet(content, metadata):
    """Process spreadsheet content to maintain cell context"""
    chunks = []
    current_sheet = ""
    
    for line in content.split('\n'):
        if line.startswith("Sheet:"):
            current_sheet = line
            chunks.append(current_sheet)
        else:
            chunks.append(f"{current_sheet} {line}")
    
    return chunks

def get_document_chunks(doc):
    """Chunk documents based on type and maintain context"""
    if doc.metadata.get('filetype', '').endswith('.xlsx'):
        return process_spreadsheet(doc.text, doc.metadata)
    else:
        # For other documents, use sentence splitter with larger chunks
        splitter = SentenceSplitter(
            chunk_size=1024,  # Larger chunks for better context
            chunk_overlap=200  # More overlap to maintain context
        )
        return splitter.split_text(doc.text)

def render_citation(citation):
    """Render a citation card with relevance-based styling"""
    score = citation.get('score', 'N/A')
    # Set default background for N/A or calculate based on score
    if score == 'N/A':
        bg_color = "hsla(200, 70%, 30%, 0.2)"  # Default blue-ish background
    else:
        score = float(score)
        hue = min(120, max(0, (score - 0.77) * 500))
        bg_color = f"hsla({hue}, 70%, 30%, 0.2)"
    
    html = f"""
    <div class="doc-card" style="background: {bg_color}; border-radius: 12px; padding: 16px; margin: 12px 0; box-shadow: 0 4px 10px rgba(0, 0, 0, 0.2);">
        <div class="doc-title" style="font-weight: bold; color: #fff;">{citation['filename']}</div>
        <div class="doc-relevance" style="color: #aaa;">Relevance Score: {citation['score']}</div>
        <div class="doc-snippet" style="color: #ddd;">{citation['snippet']}</div>
        <div class="doc-link">
            <a href="{citation['link']}" target="_blank" style="color: #4CAF50; text-decoration: none;">View Document →</a>
        </div>
    </div>
    """
    st.markdown(html, unsafe_allow_html=True)

def create_logging_sidebar():
    container = st.sidebar.container()
    container.markdown("""
        <style>
        .status-item {
            display: flex;
            align-items: center;
            margin: 8px 0;
            padding: 8px;
            background: rgba(255, 255, 255, 0.05);
            border-radius: 6px;
        }
        .status-icon {
            width: 8px;
            height: 8px;
            border-radius: 50%;
            margin-right: 12px;
        }
        .status-success { background: #4CAF50; }
        .status-error { background: #F44336; }
        .status-loading { background: #FFC107; }
        .status-text { color: #fff; }
        </style>
    """, unsafe_allow_html=True)
    
    return container

def create_header():
    st.title("🦫 Prospector Pitch Assistant")
    st.caption("An AI Assistant for Our Drive Folder")

def format_response(response_dict):
    """Format the chatbot response in a modern style"""
    citations_html = '\n'.join([
        f"""
        <div class="citation-card">
            <div class="citation-title">{citation['filename']}</div>
            <div class="citation-score">Relevance: {citation['score']}</div>
            <div class="citation-snippet">{citation['snippet']}</div>
            <a href="{citation['link']}" class="citation-link" target="_blank">View Document →</a>
        </div>
        """
        for citation in response_dict.get("primary_citations", []) + response_dict.get("secondary_citations", [])
    ])
    
    st.markdown("""
        <style>
        .response-container {
            font-family: 'Inter', sans-serif;
            background: white;
            border-radius: 12px;
            padding: 20px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        .answer {
            font-size: 16px;
            line-height: 1.6;
            color: #333;
            margin-bottom: 20px;
        }
        .citations {
            border-top: 1px solid #eee;
            padding-top: 16px;
        }
        .citation-card {
            background: #f8f9fa;
            border-radius: 8px;
            padding: 12px;
            margin: 8px 0;
        }
        .citation-title {
            font-weight: 600;
            color: #1a1a1a;
            margin-bottom: 4px;
        }
        .citation-score {
            font-size: 12px;
            color: #666;
            margin-bottom: 8px;
        }
        .citation-snippet {
            font-size: 14px;
            color: #444;
            margin-bottom: 8px;
        }
        .citation-link {
            font-size: 12px;
            color: #007bff;
        }
        </style>
        <div class="response-container">
            <div class="answer">{}</div>
            <div class="citations">
                <h3>Sources</h3>
                {}
            </div>
        </div>
    """.format(
        response_dict['answer'],
        citations_html
    ), unsafe_allow_html=True)

def create_chat_interface():
    st.markdown("""
        <style>
        .chat-container {font-family: 'Inter', sans-serif; max-width: 800px; margin: 0 auto;}
        .stTextInput > div > div > input {font-family: 'Inter', sans-serif; border-radius: 8px !important;}
        .stChatMessage {background: rgba(255, 255, 255, 0.05) !important; border-radius: 8px; margin: 8px 0;}
        .stChatInput {border-color: rgba(255, 255, 255, 0.1) !important; background: rgba(255, 255, 255, 0.05) !important;}
        </style>
    """, unsafe_allow_html=True)

def set_dark_theme():
    st.markdown("""
        <style>
        /* Dark theme base styles */
        [data-testid="stAppViewContainer"] {
            background-color: #000000;
            color: #ffffff;
        }
        [data-testid="stSidebar"] {
            background-color: #0a0a0a;
            border-right: 1px solid #1a1a1a;
        }
        .stTextInput > div > div > input {
            background-color: #1a1a1a !important;
            color: #ffffff !important;
            padding: 12px !important;
        }
        .stChatMessage {
            background-color: #1a1a1a !important;
            border: none !important;
            padding: 16px !important;
            margin: 16px 0 !important;
        }
        /* Streamlit icon modifications */
        [data-testid="stIcon"] {
            opacity: 0.7;
            transition: opacity 0.3s;
        }
        [data-testid="stIcon"]:hover {
            opacity: 1;
        }
        /* Animation for sidebar elements */
        .sidebar-container {
            animation: slideIn 0.5s ease-out;
        }
        @keyframes slideIn {
            from { transform: translateX(-100%); opacity: 0; }
            to { transform: translateX(0); opacity: 1; }
        }
        </style>
    """, unsafe_allow_html=True)

def search_web(topic):
    try:
        # Search for high-quality journalism articles
        search_results = list(search(
            f"{topic} journalism analysis report investigation",
            num_results=10
        ))
        
        # Get detailed content for each result
        detailed_results = []
        for url in search_results:
            try:
                response = requests.get(url, timeout=10)
                soup = BeautifulSoup(response.text, 'html.parser')
                
                article_text = ""
                main_content = soup.find('article') or soup.find('main') or soup.find('body')
                if main_content:
                    paragraphs = main_content.find_all('p')
                    article_text = ' '.join(p.get_text().strip() for p in paragraphs[:5])
                
                if len(article_text) > 100:
                    detailed_results.append({
                        "title": soup.title.string.strip() if soup.title else "Untitled",
                        "url": url,
                        "content": article_text
                    })
            except Exception:
                continue
        
        # Analysis prompt with specific structure
        analysis_prompt = f"""
        Analyze this journalistic topic: "{topic}"
        Provide a comprehensive analysis in exactly 5-6 sentences, covering:
        1. Current relevance and timeliness
        2. Public interest and impact
        3. Key developments or trends
        Make sure to complete all thoughts and provide a well-rounded analysis.
        """
        
        analysis = generate_openai_response(analysis_prompt)
        
        # Curation prompt with specific requirements
        if detailed_results:
            curation_prompt = f"""
            Review these articles about {topic} and provide a curated list of the 3-4 most valuable sources.
            For each article, include:
            1. The article title
            2. A one-sentence summary of its key insight
            3. The complete URL
            4. The source/publication name

            Format each entry as:
            "• [Title] - [Summary]
             Source: [Publication]
             Link: [URL]"

            Only include high-quality journalism sources with substantial coverage.
            """
            
            results_text = "\n\n".join([
                f"Article: {r['title']}\nURL: {r['url']}\nContent: {r['content'][:300]}..."
                for r in detailed_results
            ])
            
            curated_results = generate_openai_response(f"{curation_prompt}\n\nArticles to review:\n{results_text}")
        else:
            curated_results = "No relevant articles found."
        
        return analysis, curated_results
    except Exception as e:
        return f"Error: {str(e)}", "No articles found."

def generate_openai_response(prompt):
    try:
        client = openai.OpenAI(api_key=st.secrets["OPEN_AI_KEY"])
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "user", "content": prompt}
            ],
            max_tokens=500,  # Increased token limit
            temperature=0.7
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"An error occurred while generating a response: {str(e)}"

def analyze_pitch(pitch, section):
    try:
        analysis_prompt = f"""
        Analyze this pitch for the {section} section of a high school newspaper:

        Pitch: {pitch}

        Provide a detailed analysis covering:
        1. Strengths (2-3 points)
        2. Areas for Improvement (2-3 points)
        3. Specific suggestions to enhance the pitch
        4. Whether it meets the section guidelines
        5. Research suggestions to strengthen the story

        Base your analysis on these key criteria:
        - Relevance to the school community
        - Timeliness and newsworthiness
        - Uniqueness of perspective
        - Depth and complexity
        - Feasibility of reporting
        
        Format the response clearly with headers and bullet points.
        Include 2-3 specific research directions or sources to explore.
        """
        
        analysis = generate_openai_response(analysis_prompt)
        return analysis
    except Exception as e:
        return f"Error analyzing pitch: {str(e)}"

def check_originality(pitch, section):
    try:
        # Search Prospector website
        prospector_results = list(search(f"site:chsprospector.com {pitch}", num_results=5))
        issuu_results = list(search(f"site:issuu.com/prospector {pitch}", num_results=5))
        
        # Get detailed content for similarity comparison
        similar_articles = []
        
        # Process Prospector results
        for url in prospector_results + issuu_results:
            try:
                response = requests.get(url, timeout=10)
                soup = BeautifulSoup(response.text, 'html.parser')
                title = soup.title.string.strip() if soup.title else "Untitled"
                
                # Extract content
                content = ""
                if main_content := (soup.find('article') or soup.find('main')):
                    content = ' '.join(p.get_text().strip() for p in main_content.find_all('p'))
                
                if content:
                    similar_articles.append({
                        "title": title,
                        "url": url,
                        "content": content[:500]
                    })
            except Exception:
                continue
        
        # Generate originality analysis using LLM
        analysis_prompt = f"""
        Analyze this pitch for originality compared to existing articles:

        Pitch: {pitch}
        Section: {section}

        Similar articles found:
        {similar_articles}

        Provide:
        1. Originality score (0-100%)
        2. Analysis of unique elements
        3. Comparison to existing coverage
        4. Suggestions for differentiation
        
        Include specific references to similar articles if found.
        """
        
        originality_analysis = generate_openai_response(analysis_prompt)
        
        return originality_analysis, similar_articles
    except Exception as e:
        return f"Error checking originality: {str(e)}", []

def main():
    # Set up Streamlit page configuration
    st.set_page_config(
        page_title="Prospector",
        page_icon="🦫",
        layout="wide"
    )
    
    # Apply dark theme
    set_dark_theme()
    
    # Create header
    create_header()
    
    # Create chat interface styles
    create_chat_interface()
    
    # Add data source selection to main page
    if 'data_source' not in st.session_state:
        st.session_state.data_source = 'Drive'
    
    col1, col2 = st.columns([1, 2])
    with col1:
        data_source = st.selectbox("Select Data Source", ['Drive', 'Google'], key='data_source_select')
    
    with col2:
        if data_source == 'Google':
            action = st.radio("Select Action", 
                            ["Research Topic", "Critique Pitch", "Check Originality"], 
                            horizontal=True,
                            key="google_action")

    # Reset chat when data source changes
    if 'prev_data_source' not in st.session_state:
        st.session_state.prev_data_source = data_source

    if data_source != st.session_state.prev_data_source:
        st.session_state.messages = []
        st.session_state.prev_data_source = data_source
        st.rerun()
    
    try:
        # Load or create the document index with a spinner
        with st.spinner("Loading document index..."):
            if data_source == 'Drive':
                index = create_index()
                if index:
                    st.markdown('<div class="status-item"><div class="status-icon status-success"></div><div class="status-text">✓ Index Created Successfully</div></div>', unsafe_allow_html=True)
                else:
                    st.markdown('<div class="status-item"><div class="status-icon status-error"></div><div class="status-text">❌ Failed to create or load index.</div></div>', unsafe_allow_html=True)

    except Exception as e:
        st.error(f"Error creating index: {str(e)}")
        return

    # Initialize chat history in session state
    if "messages" not in st.session_state:
        st.session_state.messages = []
        st.session_state.messages.append({
            "role": "assistant",
            "content": {
                "answer": "Hello! I can help you find information in your Google Drive documents. Ask me anything about their contents!",
                "primary_citations": [],
                "secondary_citations": []
            }
        })

    # Display existing chat messages
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            if isinstance(message["content"], dict):
                st.markdown(message["content"]["answer"])
                if message["content"].get("primary_citations") or message["content"].get("secondary_citations"):
                    st.markdown("### Sources")
                    for citation in message["content"].get("primary_citations", []):
                        render_citation(citation)
                    for citation in message["content"].get("secondary_citations", []):
                        render_citation(citation)
            else:
                st.markdown(message["content"])

    # Handle user input based on data source
    if data_source == 'Drive':
        if prompt := st.chat_input("Ask me about your documents", key="drive_chat"):
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)
            
            try:
                with st.spinner("Thinking..."):
                    response = generate_response(index, prompt)
                
                if not response.get("answer"):
                    response["answer"] = "I'm sorry, I couldn't find an answer to that based on the available documents."
                
                st.session_state.messages.append({
                    "role": "assistant",
                    "content": response
                })
                
                with st.chat_message("assistant"):
                    st.markdown(response["answer"])
                    if response.get("primary_citations") or response.get("secondary_citations"):
                        st.markdown("### Sources")
                        for citation in response.get("primary_citations", []):
                            render_citation(citation)
                        for citation in response.get("secondary_citations", []):
                            render_citation(citation)
            
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")

    elif data_source == 'Google':
        if action == "Research Topic":
            if prompt := st.chat_input("Enter a topic to research", key="google_search"):
                with st.spinner("Searching and analyzing..."):
                    analysis, search_results = search_web(prompt)
                    st.write(analysis)
                    st.write("\nRelated Articles:")
                    st.write(search_results)
        
        elif action == "Check Originality":
            section = st.selectbox(
                "Select Section",
                ["News", "Opinions", "Lifestyles", "Sports", "Investigations", 
                 "Features", "Postscript", "Arts & Culture", "In-Depth", "Multimedia", "Podcast"],
                key="originality_section"
            )
            
            pitch = st.text_area("Enter your pitch", 
                               height=150,
                               placeholder="Describe your story idea in detail...",
                               key="originality_pitch")
            
            if st.button("Check Originality"):
                if not pitch:
                    st.warning("Please enter a pitch to analyze.")
                else:
                    with st.spinner("Analyzing originality..."):
                        originality_analysis, similar_articles = check_originality(pitch, section)
                        st.markdown("### Originality Analysis")
                        st.write(originality_analysis)
                        
                        if similar_articles:
                            st.markdown("### Similar Articles Found")
                            for article in similar_articles:
                                st.markdown(f"**[{article['title']}]({article['url']})**")
                                st.write(article['content'][:200] + "...")
        
        else:  # Critique Pitch
            section = st.selectbox(
                "Select Section",
                ["News", "Opinions", "Lifestyles", "Sports", "Investigations", 
                 "Features", "Postscript", "Arts & Culture", "In-Depth", "Multimedia", "Podcast"]
            )
            
            pitch = st.text_area("Enter your pitch", height=150, 
                               placeholder="Describe your story idea in detail...")
            
            if st.button("Analyze Pitch"):
                if not pitch:
                    st.warning("Please enter a pitch to analyze.")
                else:
                    with st.spinner("Analyzing pitch..."):
                        # Get pitch analysis
                        pitch_analysis = analyze_pitch(pitch, section)
                        st.markdown("### Pitch Analysis")
                        st.write(pitch_analysis)
                        
                        # Get related research
                        st.markdown("### Related Research")
                        analysis, search_results = search_web(pitch)
                        st.write(analysis)
                        st.write("\nRelevant Articles for Research:")
                        st.write(search_results)

if __name__ == "__main__":
    main()